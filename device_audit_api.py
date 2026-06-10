from fastapi import FastAPI, HTTPException, Query, UploadFile, File
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from supabase import create_client
from datetime import datetime, date
from typing import Optional, Dict, Any
from docx import Document
from pdf_report_generator import generate_executive_pdf
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import io
import logging
import os
import re
import tempfile
import shutil

try:
    from sharepoint_uploader import (
        upload_report_to_sharepoint,
        test_sharepoint_connection as sp_test_connection,
    )
    SHAREPOINT_AVAILABLE = True
except ImportError:
    SHAREPOINT_AVAILABLE = False
    logger.warning("sharepoint_uploader not available — install msal to enable SharePoint uploads")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(name)s %(levelname)s %(message)s")
logger = logging.getLogger("device_audit_api")

SUPABASE_URL = os.getenv("SUPABASE_URL", "").strip()
#SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_KEY")

SUPABASE_KEY = (
    os.getenv("SUPABASE_SERVICE_ROLE_KEY")
    or os.getenv("SUPABASE_SERVICE_KEY")
    or os.getenv("SUPABASE_KEY")
)

if not SUPABASE_URL:
    raise RuntimeError("Missing SUPABASE_URL environment variable")

if not SUPABASE_KEY:
    raise RuntimeError("Missing SUPABASE_SERVICE_KEY environment variable")

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

REPORT_STORAGE_DIR = os.getenv(
    "AUDIT_REPORT_STORAGE_DIR",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_reports"),
)

app = FastAPI(title="Device Audit API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class AuditStartRequest(BaseModel):
    site_id: Optional[str] = None
    site_name: Optional[str] = None


class AuditEntryUpdateRequest(BaseModel):
    id: Optional[str] = None
    audit_id: Optional[str] = None
    audit_batch_id: Optional[str] = None
    serial_number: Optional[str] = None
    values: Optional[Dict[str, Any]] = None


def is_true(v):
    return str(v).strip().lower() in ["true", "1", "yes", "y", "on"]


def is_false(v):
    return str(v).strip().lower() in ["false", "0", "no", "n", "off"]


def parse_bool(v, default=None):
    if v is None or v == "":
        return default
    if isinstance(v, bool):
        return v
    if is_true(v):
        return True
    if is_false(v):
        return False
    return default


def is_photo_audit_device(entry):
    text = " ".join([
        normalize_text(entry.get("device_type")),
        normalize_text(entry.get("brand_model")),
        normalize_text(entry.get("device_name")),
        normalize_text(entry.get("operating_system")),
        normalize_text(entry.get("windows_os")),
        normalize_text(entry.get("ios_version")),
    ]).lower()

    include_terms = ["ipad", "tablet", "apple device", "iphone", "ios", "samsung galaxy tab"]
    exclude_terms = ["laptop", "desktop", "windows", "phone", "voip", "printer", "whiteboard"]

    if any(term in text for term in exclude_terms):
        return False
    return any(term in text for term in include_terms)


def parse_date(v):
    if not v:
        return None
    try:
        return datetime.fromisoformat(str(v)).date()
    except Exception:
        return None


def days_old(v):
    d = parse_date(v)
    if not d:
        return None
    return (date.today() - d).days


def audit_age_days(v):
    d = parse_date(v)
    if not d:
        return None
    return (date.today() - d).days


def safe_filename(value):
    value = value or "audit-report"
    value = re.sub(r"[^a-zA-Z0-9_\-]+", "-", value)
    return value.strip("-").lower()


def site_key(value):
    return str(value or "").strip().lower()


def normalize_text(value):
    return str(value or "").strip()


def resolve_site(site_id=None, site_name=None):
    if site_id:
        r = supabase.table("sites").select("*").eq("site_id", site_id).limit(1).execute()
        if r.data:
            return r.data[0]

    if site_name:
        r = supabase.table("sites").select("*").eq("site_name", site_name).limit(1).execute()
        if r.data:
            return r.data[0]

    raise HTTPException(status_code=404, detail="Site not found")


def clean_entry_payload(payload):
    allowed = {
        "security_settings_applied",
        "onedrive_status",
        "windows_updates",
        "security_check",
        "windows_os",
        "ios_version",
        "update_status",
        "camera_sync_off",
        "onedrive_sync_on",
        "photos_count",
        "photos_date",
        "total_photos",
        "ignore_flag",
        "notes",
        "site_name",
        "assigned_user_room",
        "device_id",
        "device_present",
        "device_name",
        "device_type",
        "assigned_user",
        "assigned_room",
        "os_version",
        "ip_address",
        "mac_address_wifi",
        "mac_address_lan",
        "password_issue",
        "weak_pin_detected",
        "security_breach",
        "breach_notes",
        "recommended_action",
    }
    return {k: v for k, v in payload.items() if k in allowed}


def clean_device_payload(payload):
    allowed = {
        "asset_tag",
        "device_type",
        "brand_model",
        "device_name",
        "date_of_purchase",
        "warranty_expiry",
        "assigned_user_room",
        "assigned_user",
        "assigned_room",
        "status",
        "intended_use",
        "connectivity",
        "recording_storage_capabilities",
        "ip_address",
        "mac_address_wifi",
        "mac_address_lan",
        "processor",
        "ram",
        "storage",
        "operating_system",
        "os_version",
        "notes",
        "is_active",
    }
    return {k: v for k, v in payload.items() if k in allowed}



def get_site_thresholds(site_id):
    """Return photo/audit thresholds and policy flags for a site.

    Falls back to safe defaults if the site cannot be resolved.
    """
    defaults = {
        "photo_retention_days": 31,
        "photo_count_threshold": 1000,
        "audit_frequency_days": 90,
        "consent_forms_current": None,
        "photo_policy_documented": None,
        "staff_training_current": None,
        "encryption_enforced": None,
        "access_controls_documented": None,
        "data_retention_policy": None,
        "incident_response_plan": None,
        "privacy_impact_assessment": None,
        "third_party_data_agreement": None,
    }
    if not site_id:
        return defaults
    try:
        r = supabase.table("sites").select(
            "photo_retention_days,photo_count_threshold,audit_frequency_days,"
            "consent_forms_current,photo_policy_documented,staff_training_current,"
            "encryption_enforced,access_controls_documented,data_retention_policy,"
            "incident_response_plan,privacy_impact_assessment,third_party_data_agreement"
        ).eq("site_id", site_id).limit(1).execute()
        if r.data:
            site = r.data[0]
            for k, v in site.items():
                if v is not None:
                    defaults[k] = v
    except Exception:
        pass
    return defaults


def classify(entry, retention_days=31, photo_threshold=1000):
    if is_true(entry.get("ignore_flag")):
        return "ignored", ["Ignored"]

    critical = []
    warning = []

    device_present = parse_bool(entry.get("device_present"), default=True)

    if device_present is False:
        critical.append("Missing device")
        return "critical", critical

    if (
        parse_bool(entry.get("password_issue"), default=False) is True
        or parse_bool(entry.get("weak_pin_detected"), default=False) is True
        or parse_bool(entry.get("security_breach"), default=False) is True
    ):
        critical.append("Security issue")

    photo_device = is_photo_audit_device(entry)

    if photo_device:
        photo_age = days_old(entry.get("photos_date"))
        if photo_age is not None and photo_age > retention_days:
            critical.append(f"Photos older than {retention_days} days")

        try:
            photos = int(entry.get("photos_count") or entry.get("total_photos") or 0)
        except Exception:
            photos = 0

        if photos > photo_threshold:
            warning.append("Photo count over threshold")

        camera_sync_off = parse_bool(entry.get("camera_sync_off"), default=None)
        if camera_sync_off is False:
            warning.append("Camera sync not off")

        onedrive_sync_on = parse_bool(entry.get("onedrive_sync_on"), default=None)
        if onedrive_sync_on is False:
            warning.append("OneDrive sync not confirmed")

    if entry.get("update_status") and str(entry.get("update_status")).strip().lower() not in ["ok", "complete", "updated", "current", "yes", "done"]:
        warning.append("Update status requires review")

    if critical:
        return "critical", critical

    if warning:
        return "warning", warning

    return "ok", ["OK"]

def build_rows(entries, site_id=None):
    rows = []
    thresholds = get_site_thresholds(site_id)
    retention_days  = int(thresholds.get("photo_retention_days")  or 31)
    photo_threshold = int(thresholds.get("photo_count_threshold") or 1000)
    for entry in entries:
        status, reasons = classify(entry, retention_days=retention_days, photo_threshold=photo_threshold)
        row = dict(entry)
        row["status"] = status
        row["reasons"] = reasons
        rows.append(row)
    return rows


def get_audit_batch(audit_batch_id: str):
    r = (
        supabase.table("device_audits")
        .select("*")
        .eq("audit_batch_id", audit_batch_id)
        .limit(1)
        .execute()
    )
    if not r.data:
        raise HTTPException(status_code=404, detail="Audit batch not found")
    return r.data[0]


def get_entries_for_batch(audit_batch_id: str):
    entries = (
        supabase.table("audit_entries")
        .select("*")
        .eq("audit_batch_id", audit_batch_id)
        .execute()
        .data
        or []
    )
    # Look up site_id from the audit batch for per-site thresholds
    batch_rec = (
        supabase.table("device_audits")
        .select("site_id")
        .eq("audit_batch_id", audit_batch_id)
        .limit(1)
        .execute()
    )
    _site_id = batch_rec.data[0].get("site_id") if batch_rec.data else None
    rows = build_rows(entries, site_id=_site_id)
    return sorted(
        rows,
        key=lambda r: (
            room_name(r).lower(),
            normalize_text(r.get("device_name") or r.get("brand_model") or r.get("device_type")).lower(),
            normalize_text(r.get("serial_number")).lower(),
        ),
    )


def build_report(audit_batch_id: str):
    batch = get_audit_batch(audit_batch_id)
    rows = get_entries_for_batch(audit_batch_id)

    site_config = get_site_thresholds(batch.get("site_id"))
    retention_days = int(site_config.get("photo_retention_days") or 31)
    photo_threshold = int(site_config.get("photo_count_threshold") or 1000)

    active_rows = [r for r in rows if not is_true(r.get("ignore_flag"))]

    critical = [r for r in active_rows if r.get("status") == "critical"]
    warnings = [r for r in active_rows if r.get("status") == "warning"]
    ok = [r for r in active_rows if r.get("status") == "ok"]

    missing = [r for r in active_rows if r.get("device_present") is False]

    photo_breaches = []
    for r in active_rows:
        photo_age = days_old(r.get("photos_date"))
        if photo_age is not None and photo_age > retention_days:
            photo_breaches.append(r)

    photo_warnings = []
    for r in active_rows:
        try:
            photos = int(r.get("photos_count") or r.get("total_photos") or 0)
        except Exception:
            photos = 0
        if photos > photo_threshold:
            photo_warnings.append(r)

    security_issues = [
        r for r in active_rows
        if r.get("password_issue") is True
        or r.get("weak_pin_detected") is True
        or r.get("security_breach") is True
    ]

    if critical:
        report_status = "Critical"
    elif warnings:
        report_status = "Warning"
    else:
        report_status = "Compliant"

    recommendations = []

    if missing:
        recommendations.append("Immediately locate or formally retire all devices marked as missing.")
    if photo_breaches:
        recommendations.append(f"Remove or archive photos older than {retention_days} days in line with photo retention expectations.")
    if security_issues:
        recommendations.append("Reset weak/default passwords and PINs and confirm security settings are applied.")
    if warnings:
        recommendations.append("Review warning items and record corrective actions against each device.")
    if not recommendations:
        recommendations.append("Continue routine audit monitoring and retain this report as compliance evidence.")

    summary = {
        "total_devices": len(active_rows),
        "ignored": len(rows) - len(active_rows),
        "critical": len(critical),
        "warnings": len(warnings),
        "ok": len(ok),
        "missing": len(missing),
        "photo_retention_breaches": len(photo_breaches),
        "photo_count_warnings": len(photo_warnings),
        "security_issues": len(security_issues),
    }

    return {
        "audit_batch_id": audit_batch_id,
        "site_name": batch.get("site_name"),
        "site_id": batch.get("site_id"),
        "site_config": site_config,
        "audit_date": batch.get("audit_date"),
        "generated_at": datetime.now().isoformat(),
        "overall_status": report_status,
        "summary": summary,
        "critical_incidents": critical,
        "warnings": warnings,
        "missing_devices": missing,
        "photo_retention_breaches": photo_breaches,
        "security_issues": security_issues,
        "recommendations": recommendations,
        "rows": active_rows,
        "report_generated": batch.get("report_generated"),
        "report_url": batch.get("report_url"),
    }


def calculate_score(summary):
    score = 100
    score -= int(summary.get("critical") or 0) * 20
    score -= int(summary.get("missing") or 0) * 15
    score -= int(summary.get("security_issues") or 0) * 12
    score -= int(summary.get("photo_retention_breaches") or 0) * 10
    score -= int(summary.get("warnings") or 0) * 5
    return max(0, min(100, score))


def calculate_risk(summary):
    if (
        int(summary.get("critical") or 0) > 0
        or int(summary.get("missing") or 0) > 0
        or int(summary.get("security_issues") or 0) > 0
    ):
        return "High Risk"

    if (
        int(summary.get("warnings") or 0) > 0
        or int(summary.get("photo_retention_breaches") or 0) > 0
    ):
        return "Needs Attention"

    return "On Track"


def decorate_director_report(report):
    summary = report.get("summary") or {}
    risk = calculate_risk(summary)
    score = calculate_score(summary)
    age = audit_age_days(report.get("audit_date"))

    return {
        "site_id": report.get("site_id"),
        "site_name": report.get("site_name"),
        "audit_batch_id": report.get("audit_batch_id"),
        "latest_batch_id": report.get("audit_batch_id"),
        "audit_date": report.get("audit_date"),
        "audit_age_days": age,
        "stale_audit": False,
        "risk_label": risk,
        "overall_status": "Compliant" if risk == "On Track" else "Not Compliant",
        "compliance_score": score,
        "report_generated": is_true(report.get("report_generated")) or bool(report.get("report_url")),
        "report_url": report.get("report_url"),
        "summary": summary,
    }


def format_au_date(value):
    if value is None or value == "":
        return None

    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")

    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")

    text = str(value).strip()
    if not text:
        return None

    # Handles ISO dates and datetimes from Supabase, for example:
    # 2026-05-29, 2026-05-29T06:15:58+00:00, 2026-05-29 06:15:58+00
    try:
        clean = text.replace("Z", "+00:00")
        if "T" in clean or " " in clean:
            return datetime.fromisoformat(clean).strftime("%d/%m/%Y")
        return datetime.fromisoformat(clean).date().strftime("%d/%m/%Y")
    except Exception:
        pass

    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text[:10], fmt).strftime("%d/%m/%Y")
        except Exception:
            continue

    return None


def format_report_value(value):
    if value is None or value == "":
        return "N/A"
    if isinstance(value, bool):
        return "Yes" if value else "No"

    au_date = format_au_date(value)
    if au_date and re.match(r"^\d{4}-\d{2}-\d{2}", str(value).strip()):
        return au_date

    return str(value)


def room_name(row):
    return (
        row.get("assigned_user_room")
        or row.get("assigned_room")
        or row.get("assigned_user")
        or "Unassigned"
    )


def device_label(row):
    return row.get("brand_model") or row.get("device_name") or row.get("device_type") or "Device"


def photo_count(row):
    try:
        return int(row.get("photos_count") or row.get("total_photos") or 0)
    except Exception:
        return 0


def is_missing(row):
    return parse_bool(row.get("device_present"), default=True) is False


def is_photo_action(row):
    if not is_photo_audit_device(row):
        return False
    photo_age = days_old(row.get("photos_date"))
    return photo_count(row) > 1000 or (photo_age is not None and photo_age > 31)


def has_credential_issue(row):
    text = " ".join([
        normalize_text(row.get("notes")),
        normalize_text(row.get("breach_notes")),
        normalize_text(row.get("recommended_action")),
    ]).lower()
    return (
        row.get("password_issue") is True
        or row.get("weak_pin_detected") is True
        or row.get("security_breach") is True
        or "password" in text
        or "pin" in text
        or "123456" in text
        or "sticker" in text
        or "credential" in text
    )


def has_hardware_issue(row):
    text = " ".join([
        normalize_text(row.get("notes")),
        normalize_text(row.get("breach_notes")),
        normalize_text(row.get("recommended_action")),
    ]).lower()
    keywords = ["broken", "battery", "hinge", "keyboard", "screen", "touch", "protector", "fault", "repair", "flat", "missing key"]
    return any(k in text for k in keywords)


def row_status_text(row):
    if is_missing(row):
        return "MISSING - not sighted during audit"

    parts = ["Active"]

    if is_photo_action(row):
        parts.append("photo reduction required")
    if has_credential_issue(row):
        parts.append("credential remediation required")
    if has_hardware_issue(row):
        parts.append("hardware remediation required")
    if row.get("status") == "warning" and len(parts) == 1:
        parts.append("review required")
    if row.get("status") == "critical" and len(parts) == 1:
        parts.append("critical review required")

    return " / ".join(parts)


def add_table(doc, title, rows, columns):
    doc.add_heading(title, level=2)

    if not rows:
        doc.add_paragraph("None recorded.")
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, (_, label) in enumerate(columns):
        header_cells[i].text = label

    for row in rows:
        cells = table.add_row().cells
        for i, (key, _) in enumerate(columns):
            value = row.get(key)
            if isinstance(value, list):
                value = ", ".join(value)
            cells[i].text = format_report_value(value)


def add_kv_table(doc, pairs):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Count"

    for label, value in pairs:
        cells = table.add_row().cells
        cells[0].text = format_report_value(label)
        cells[1].text = format_report_value(value)


def add_detailed_room_table(doc, room, rows):
    doc.add_heading(room, level=2)

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = [
        "Device Type",
        "Brand / Model",
        "Serial",
        "Windows OS / iOS Update",
        "OneDrive / Camera / iCloud Sync",
        "Photos",
        "Issues / Notes",
        "Status",
    ]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = format_report_value(row.get("device_type"))
        cells[1].text = format_report_value(row.get("brand_model") or row.get("device_name"))
        cells[2].text = format_report_value(row.get("serial_number"))
        cells[3].text = (
            f"Windows OS: {format_report_value(row.get('windows_os') or row.get('os_version'))}; "
            f"Windows updates: {format_report_value(row.get('windows_updates'))}; "
            f"Security check: {format_report_value(row.get('security_check'))}; "
            f"iOS update: {format_report_value(row.get('ios_version') or row.get('update_status'))}"
        )
        cells[4].text = (
            f"OneDrive: {format_report_value(row.get('onedrive_status'))}; "
            f"Camera sync off: {format_report_value(row.get('camera_sync_off'))}; "
            f"OneDrive sync on: {format_report_value(row.get('onedrive_sync_on'))}; "
            f"iCloud sync off: {format_report_value(row.get('icloud_sync_off'))}"
        )
        cells[5].text = f"{photo_count(row)} photos; Date: {format_report_value(row.get('photos_date'))}"
        cells[6].text = (
            f"Issues: {format_report_value(row.get('breach_notes') or row.get('recommended_action'))}; "
            f"Notes: {format_report_value(row.get('notes'))}"
        )
        cells[7].text = row_status_text(row)


def build_critical_actions(rows):
    actions = []

    for row in rows:
        room = room_name(row)
        serial = row.get("serial_number")
        model = device_label(row)
        dtype = row.get("device_type") or "Device"
        notes = normalize_text(row.get("notes"))
        breach = normalize_text(row.get("breach_notes"))
        issue_text = breach or notes

        if is_missing(row):
            actions.append({
                "incident_type": "Missing Asset",
                "room": room,
                "serial": serial,
                "issue": f"{dtype} {model} was missing from audit",
                "action": "Locate, verify and disable access if unrecovered.",
            })

        if is_photo_action(row):
            actions.append({
                "incident_type": "Photo Retention",
                "room": room,
                "serial": serial,
                "issue": f"{photo_count(row)} photos recorded on {format_report_value(row.get('photos_date'))}",
                "action": "Reduce retained photos to under 30 days after confirming approved upload/storage.",
            })

        if has_credential_issue(row):
            actions.append({
                "incident_type": "Credential Security",
                "room": room,
                "serial": serial,
                "issue": issue_text or "Credential or PIN concern recorded",
                "action": "Remove visible/shared credentials and apply unique device/room credentials.",
            })

        if has_hardware_issue(row):
            actions.append({
                "incident_type": "Hardware",
                "room": room,
                "serial": serial,
                "issue": issue_text or "Hardware issue recorded",
                "action": "Repair, replace or formally risk accept the device before continued operational use.",
            })

    return actions


def set_document_layout(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)


def create_docx_report(report):
    doc = Document()
    set_document_layout(doc)

    rows = report.get("rows", [])
    active_rows = [r for r in rows if not is_missing(r)]
    missing_rows = [r for r in rows if is_missing(r)]
    photo_rows = [r for r in active_rows if is_photo_action(r)]
    credential_rows = [r for r in active_rows if has_credential_issue(r)]
    hardware_rows = [r for r in active_rows if has_hardware_issue(r)]
    critical_actions = build_critical_actions(rows)

    device_type_counts = {}
    for row in rows:
        dtype = format_report_value(row.get("device_type"))
        device_type_counts[dtype] = device_type_counts.get(dtype, 0) + 1

    title = doc.add_heading(f"Site Audit Report: {format_report_value(report.get('site_name'))}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Digital Device, Photo Retention & Cloud Sync Compliance Audit")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Report Date: {format_report_value(report.get('audit_date'))}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Date: {format_report_value(report.get('audit_date'))}")

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        f"Assessment compiled from {format_report_value(report.get('site_name'))} device audit data. "
        f"This report includes all {len(rows)} active audit records for the selected batch, excluding rows marked ignored. "
        "Devices are grouped by room/user and reported with available Windows checks, iOS checks, camera sync, "
        "OneDrive sync, photo counts, device issues and notes."
    )

    doc.add_heading("Key Findings", level=2)
    key_findings = [
        f"{len(rows)} total device records listed in the audit dataset.",
        f"{len(missing_rows)} devices were identified as missing from audit and escalated to the Missing Assets Register.",
        f"{len(photo_rows)} active devices have recorded photo counts or photo retention indicators requiring review.",
        f"{len(credential_rows)} active devices have credential, password or PIN remediation indicators.",
        f"{len(hardware_rows)} active devices have hardware remediation indicators.",
    ]
    for finding in key_findings:
        doc.add_paragraph(finding, style="List Bullet")

    metric_pairs = [
        ("Total listed records", len(rows)),
        ("Active / sighted records", len(active_rows)),
        ("Missing from audit", len(missing_rows)),
        ("Photo retention action items", len(photo_rows)),
        ("Credential remediation items", len(credential_rows)),
        ("Hardware remediation items", len(hardware_rows)),
    ]
    for dtype, count in sorted(device_type_counts.items()):
        metric_pairs.append((dtype, count))

    add_kv_table(doc, metric_pairs)

    doc.add_heading("Priority Actions", level=2)
    priority_actions = []
    if missing_rows:
        priority_actions.append("Investigate and locate all missing devices immediately.")
    if photo_rows:
        priority_actions.append("Reduce retained photos to under 30 days where applicable.")
    if credential_rows:
        priority_actions.append("Remove visible/shared credentials and replace standard PINs with unique room/device credentials.")
    if hardware_rows:
        priority_actions.append("Repair, replace or risk accept faulty devices that affect secure and reliable operation.")
    priority_actions.append("Validate incomplete audit rows where Windows/iOS/sync data was not recorded.")

    for action in priority_actions:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("Critical Incident: Missing Devices", level=1)
    doc.add_paragraph(
        "The following devices were identified as missing from audit. These are not to be treated as normal active records. "
        "They require physical verification, access review and, where child data could be present, privacy/security escalation."
    )

    add_table(
        doc,
        "Missing Assets Requiring Immediate Action",
        [{
            "device_type": r.get("device_type"),
            "assigned_user_room": room_name(r),
            "brand_model": r.get("brand_model") or r.get("device_name"),
            "serial_number": r.get("serial_number"),
            "risk_note": "Potential data or operational exposure",
            "required_action": "Locate / verify / disable access if unrecovered",
        } for r in missing_rows],
        [
            ("device_type", "Device Type"),
            ("assigned_user_room", "Assigned User / Room"),
            ("brand_model", "Brand / Model"),
            ("serial_number", "Serial Number"),
            ("risk_note", "Risk Note"),
            ("required_action", "Required Action"),
        ],
    )

    doc.add_heading("Detailed Audit by Room / User", level=1)
    grouped = {}
    for row in sorted(rows, key=lambda r: (room_name(r).lower(), normalize_text(r.get("device_type")).lower(), normalize_text(r.get("device_name")).lower())):
        grouped.setdefault(room_name(row), []).append(row)

    for room, room_rows in grouped.items():
        add_detailed_room_table(doc, room, room_rows)

    add_table(
        doc,
        "Photo Retention Review",
        [{
            "assigned_user_room": room_name(r),
            "device": r.get("brand_model") or r.get("device_name"),
            "serial_number": r.get("serial_number"),
            "photo_count": photo_count(r),
            "photos_date": r.get("photos_date"),
            "onedrive_sync_on": r.get("onedrive_sync_on"),
            "camera_sync_off": r.get("camera_sync_off"),
            "icloud_sync_off": r.get("icloud_sync_off"),
            "action_required": "Reduce to under 30 days",
        } for r in sorted(photo_rows, key=lambda x: photo_count(x), reverse=True)],
        [
            ("assigned_user_room", "Room/User"),
            ("device", "Device"),
            ("serial_number", "Serial"),
            ("photo_count", "Photo Count"),
            ("photos_date", "Photo Date"),
            ("onedrive_sync_on", "OneDrive Sync"),
            ("camera_sync_off", "Camera Sync Off"),
            ("icloud_sync_off", "iCloud Sync Off"),
            ("action_required", "Action Required"),
        ],
    )

    add_table(
        doc,
        "Missing Assets Register",
        [{
            "device_type": r.get("device_type"),
            "assigned_user_room": room_name(r),
            "brand_model": r.get("brand_model") or r.get("device_name"),
            "serial_number": r.get("serial_number"),
            "risk_note": "Missing from audit",
            "required_action": "Locate / verify / disable access if unrecovered",
        } for r in missing_rows],
        [
            ("device_type", "Device Type"),
            ("assigned_user_room", "Assigned User / Room"),
            ("brand_model", "Brand / Model"),
            ("serial_number", "Serial Number"),
            ("risk_note", "Risk Note"),
            ("required_action", "Required Action"),
        ],
    )

    add_table(
        doc,
        "Critical Incidents & Immediate Actions",
        critical_actions,
        [
            ("incident_type", "Incident Type"),
            ("room", "Room/User"),
            ("serial", "Serial"),
            ("issue", "Issue"),
            ("action", "Action"),
        ],
    )

    filename = f"{safe_filename(report.get('site_name'))}-{report.get('audit_batch_id')}.docx"
    path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(path)

    return path, filename


def ensure_report_storage_dir():
    os.makedirs(REPORT_STORAGE_DIR, exist_ok=True)
    return REPORT_STORAGE_DIR


def stored_report_filename(report):
    return f"{safe_filename(report.get('site_name'))}-{report.get('audit_batch_id')}.docx"


def existing_report_path_from_batch(batch):
    report_url = batch.get("report_url")
    if report_url and os.path.exists(report_url):
        return report_url
    return None


def mark_report_stored(audit_batch_id: str, report_path: str):
    payload = {
        "report_generated": True,
        "report_url": report_path,
    }

    result = (
        supabase.table("device_audits")
        .update(payload)
        .eq("audit_batch_id", audit_batch_id)
        .execute()
    )

    if not result.data:
        raise HTTPException(status_code=500, detail="Report generated but audit batch metadata update failed")

    return result.data[0]


def invalidate_stored_report(audit_batch_id: Optional[str]):
    if not audit_batch_id:
        return
    try:
        supabase.table("device_audits").update({
            "report_generated": False,
            "report_url": None,
        }).eq("audit_batch_id", audit_batch_id).execute()
    except Exception:
        pass


def generate_and_store_report(audit_batch_id: str, force: bool = False):
    batch = get_audit_batch(audit_batch_id)

    if not force:
        existing_path = existing_report_path_from_batch(batch)
        if existing_path:
            return existing_path, os.path.basename(existing_path), False

    report = build_report(audit_batch_id)
    temp_path, _ = create_docx_report(report)

    ensure_report_storage_dir()
    filename = stored_report_filename(report)
    final_path = os.path.join(REPORT_STORAGE_DIR, filename)

    if os.path.abspath(temp_path) != os.path.abspath(final_path):
        shutil.copyfile(temp_path, final_path)

    mark_report_stored(audit_batch_id, final_path)

    return final_path, filename, True


def active_flag_value(value):
    if value is None:
        return True
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() not in ["false", "0", "no", "n", "inactive"]


def is_active_device(device):
    return active_flag_value(device.get("is_active"))


def device_type_in_register(value):
    text = normalize_text(value).lower()
    if not text:
        return True
    register_terms = [
        "tablet",
        "ipad",
        "laptop",
        "notebook",
        "phone",
        "iphone",
        "mobile",
        "handset",
    ]
    return any(term in text for term in register_terms)


def latest_audit_for_site(site_id):
    audits = supabase.table("device_audits").select("*").eq("site_id", site_id).execute().data or []
    audits = sorted(audits, key=lambda a: str(a.get("audit_date") or ""), reverse=True)
    return audits[0] if audits else None


def latest_entries_by_device(audit_batch_id):
    if not audit_batch_id:
        return {}

    entries = supabase.table("audit_entries").select("*").eq("audit_batch_id", audit_batch_id).execute().data or []
    indexed = {}

    for entry in entries:
        status, reasons = classify(entry)
        row = dict(entry)
        row["status"] = status
        row["reasons"] = reasons

        device_id = row.get("device_id")
        serial_number = row.get("serial_number")

        if device_id:
            indexed[f"device_id:{device_id}"] = row
        if serial_number:
            indexed[f"serial:{serial_number}"] = row

    return indexed


def active_device_register_rows(site_id):
    latest_audit = latest_audit_for_site(site_id)
    latest_batch_id = latest_audit.get("audit_batch_id") if latest_audit else None
    latest_audit_date = latest_audit.get("audit_date") if latest_audit else None
    audit_entries = latest_entries_by_device(latest_batch_id)

    devices = supabase.table("devices").select("*").eq("site_id", site_id).execute().data or []
    rows = []

    for device in devices:
        if not is_active_device(device):
            continue

        if not device_type_in_register(device.get("device_type")):
            continue

        entry = None
        device_id = device.get("device_id")
        serial_number = device.get("serial_number")

        if device_id:
            entry = audit_entries.get(f"device_id:{device_id}")
        if not entry and serial_number:
            entry = audit_entries.get(f"serial:{serial_number}")

        audit_status = entry.get("status") if entry else ""
        audit_reasons = entry.get("reasons") if entry else []

        rows.append({
            "device_id": device.get("device_id"),
            "site_id": device.get("site_id"),
            "serial_number": device.get("serial_number"),
            "asset_tag": device.get("asset_tag"),
            "device_type": device.get("device_type"),
            "brand_model": device.get("brand_model"),
            "device_name": device.get("device_name"),
            "assigned_user_room": device.get("assigned_user_room") or device.get("assigned_room") or device.get("assigned_user"),
            "assigned_user": device.get("assigned_user"),
            "assigned_room": device.get("assigned_room"),
            "operating_system": device.get("operating_system"),
            "os_version": device.get("os_version"),
            "ip_address": device.get("ip_address"),
            "status": device.get("status"),
            "is_active": is_active_device(device),
            "latest_audit_batch_id": latest_batch_id,
            "latest_audit_date": latest_audit_date,
            "latest_audit_status": audit_status,
            "latest_audit_reasons": audit_reasons,
            "compliance": "Compliant" if audit_status == "ok" else "Not Compliant",
            "notes": device.get("notes"),
        })

    rows = sorted(
        rows,
        key=lambda r: (
            str(r.get("device_type") or "").lower(),
            str(r.get("assigned_user_room") or "").lower(),
            str(r.get("device_name") or "").lower(),
            str(r.get("serial_number") or "").lower(),
        ),
    )

    return rows, latest_audit


@app.get("/")
def root():
    return {"status": "ok"}


@app.get("/debug/routes")
def debug_routes():
    return [r.path for r in app.routes]


@app.get("/sites")
def get_sites():
    r = supabase.table("sites").select("site_id, site_name").order("site_name").execute()
    return r.data or []


@app.get("/audit-batches")
def get_audit_batches(site_id: Optional[str] = Query(default=None)):
    q = supabase.table("device_audits").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    try:
        q = q.order("audit_date", desc=True)
    except TypeError:
        q = q.order("audit_date")
    return q.execute().data or []


@app.get("/devices")
def get_devices(site_id: Optional[str] = Query(default=None), active: Optional[str] = Query(default=None)):
    q = supabase.table("devices").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    devices = q.execute().data or []

    if active is not None and str(active).lower() not in ["all", ""]:
        want_active = str(active).lower() in ["true", "1", "yes", "y", "active"]
        devices = [d for d in devices if is_active_device(d) is want_active]

    return devices


@app.get("/active-device-register")
def get_active_device_register(site_id: str = Query(...)):
    site = resolve_site(site_id=site_id)
    rows, latest_audit = active_device_register_rows(site["site_id"])

    return {
        "report_type": "Active Device Register",
        "report_date": date.today().isoformat(),
        "site_id": site.get("site_id"),
        "site_name": site.get("site_name"),
        "latest_audit_batch_id": latest_audit.get("audit_batch_id") if latest_audit else None,
        "latest_audit_date": latest_audit.get("audit_date") if latest_audit else None,
        "summary": {
            "active_devices": len(rows),
        },
        "rows": rows,
    }


# ── CSV Export ────────────────────────────────────────────────────────────────
@app.get("/devices/export.csv")
def export_devices_csv(site_id: Optional[str] = Query(default=None)):
    q = supabase.table("devices").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    devices = q.execute().data or []

    fields = ['device_id', 'site_id', 'site_name', 'asset_tag', 'device_type', 'brand_model', 'device_name', 'date_of_purchase', 'warranty_expiry', 'assigned_user_room', 'assigned_user', 'assigned_room', 'status', 'intended_use', 'connectivity', 'recording_storage_capabilities', 'ip_address', 'mac_address_wifi', 'mac_address_lan', 'processor', 'ram', 'storage', 'operating_system', 'os_version', 'notes', 'is_active']
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=fields, extrasaction="ignore")
    writer.writeheader()
    for d in devices:
        writer.writerow({f: d.get(f, "") for f in fields})

    filename = f"devices-{site_id or 'all'}.csv"
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── CSV Import ────────────────────────────────────────────────────────────────
@app.post("/devices/import")
async def import_devices_csv(file: UploadFile = File(...)):
    content = await file.read()
    try:
        text = content.decode("utf-8-sig")  # handle BOM
    except Exception:
        raise HTTPException(status_code=400, detail="Could not decode file as UTF-8.")

    reader = csv.DictReader(io.StringIO(text))
    created = updated = errors = 0
    error_rows = []

    for i, row in enumerate(reader, start=2):
        site_id = (row.get("site_id") or "").strip()
        if not site_id:
            errors += 1
            error_rows.append({"row": i, "error": "Missing site_id"})
            continue

        data = clean_device_payload({k: (v.strip() if isinstance(v, str) else v) for k, v in row.items()})
        # Convert is_active string to bool
        if "is_active" in data:
            data["is_active"] = str(data["is_active"]).lower() in ("true", "1", "yes", "y")

        device_id = (row.get("device_id") or "").strip()
        try:
            if device_id:
                # Update existing
                r = supabase.table("devices").update(data).eq("device_id", device_id).execute()
                if r.data:
                    updated += 1
                else:
                    errors += 1
                    error_rows.append({"row": i, "error": f"device_id {device_id} not found"})
            else:
                # Create new
                data["site_id"] = site_id
                site_name = (row.get("site_name") or "").strip()
                if site_name:
                    data["site_name"] = site_name
                r = supabase.table("devices").insert(data).execute()
                if r.data:
                    created += 1
                else:
                    errors += 1
                    error_rows.append({"row": i, "error": "Insert failed"})
        except Exception as e:
            errors += 1
            error_rows.append({"row": i, "error": str(e)})

    return {
        "created": created,
        "updated": updated,
        "errors": errors,
        "error_details": error_rows[:20]
    }


# ── Device Create ─────────────────────────────────────────────────────────────
@app.post("/devices/create")
def create_device(payload: Dict[str, Any]):
    site_id = (payload.get("site_id") or "").strip()
    if not site_id:
        raise HTTPException(status_code=400, detail="site_id is required")
    data = clean_device_payload(payload)
    data["site_id"] = site_id
    site_name = (payload.get("site_name") or "").strip()
    if site_name:
        data["site_name"] = site_name
    r = supabase.table("devices").insert(data).execute()
    if not r.data:
        raise HTTPException(status_code=500, detail="Failed to create device")
    return {"success": True, "device": r.data[0]}


@app.get("/devices/{device_id}")
def get_device(device_id: str):
    r = (
        supabase.table("devices")
        .select("*")
        .eq("device_id", device_id)
        .limit(1)
        .execute()
    )

    if not r.data:
        raise HTTPException(status_code=404, detail="Device not found")

    return r.data[0]


@app.post("/devices/{device_id}/update")
def update_device(device_id: str, payload: Dict[str, Any]):
    data = clean_device_payload(payload)

    result = (
        supabase.table("devices")
        .update(data)
        .eq("device_id", device_id)
        .execute()
    )

    if not result.data:
        raise HTTPException(status_code=404, detail="Update failed")

    return {"success": True, "device": result.data[0]}


@app.post("/audits/start")
def start_audit(req: AuditStartRequest):
    site = resolve_site(req.site_id, req.site_name)

    audit_payload = {
        "site_id": site["site_id"],
        "site_name": site["site_name"],
        "audit_date": date.today().isoformat(),
        "audit_type": "Device Audit",
        "report_generated": False,
    }

    audit = supabase.table("device_audits").insert(audit_payload).execute()

    if not audit.data:
        raise HTTPException(status_code=500, detail="Audit create failed")

    batch = audit.data[0]
    batch_id = batch.get("audit_batch_id")

    if not batch_id:
        raise HTTPException(status_code=500, detail=f"No audit_batch_id returned: {batch}")

    devices = (
        supabase.table("devices")
        .select("*")
        .eq("site_id", site["site_id"])
        .execute()
        .data
        or []
    )

    entries = []

    for d in devices:
        entries.append({
            "audit_batch_id": batch_id,
            "site_name": site["site_name"],
            "serial_number": d.get("serial_number"),
            "audit_date": datetime.now().isoformat(),
            "device_id": d.get("device_id"),
            "device_present": True,
            "device_name": d.get("device_name"),
            "device_type": d.get("device_type"),
            "assigned_user": d.get("assigned_user"),
            "assigned_room": d.get("assigned_room"),
            "assigned_user_room": d.get("assigned_user_room") or d.get("assigned_room") or d.get("assigned_user"),
            "windows_os": d.get("operating_system"),
            "os_version": d.get("os_version"),
            "ios_version": None,
            "ip_address": d.get("ip_address"),
            "mac_address_wifi": d.get("mac_address_wifi"),
            "mac_address_lan": d.get("mac_address_lan"),
            "camera_sync_off": None,
            "onedrive_sync_on": None,
            "photos_count": None,
            "total_photos": None,
            "photos_date": None,
            "password_issue": False,
            "weak_pin_detected": False,
            "security_breach": False,
            "breach_notes": None,
            "recommended_action": None,
            "notes": d.get("notes"),
            "ignore_flag": False,
        })

    if entries:
        supabase.table("audit_entries").insert(entries).execute()

    created = (
        supabase.table("audit_entries")
        .select("*")
        .eq("audit_batch_id", batch_id)
        .execute()
        .data
        or []
    )

    return {
        "audit_batch_id": batch_id,
        "entries": build_rows(created, site_id=site["site_id"]),
    }


@app.get("/audit-entries/{audit_batch_id}")
def get_audit_entries(audit_batch_id: str):
    entries = (
        supabase.table("audit_entries")
        .select("*")
        .eq("audit_batch_id", audit_batch_id)
        .execute()
        .data
        or []
    )

    # Look up site_id for per-site classification thresholds
    _batch = (
        supabase.table("device_audits")
        .select("site_id")
        .eq("audit_batch_id", audit_batch_id)
        .limit(1)
        .execute()
    )
    _site_id = _batch.data[0].get("site_id") if _batch.data else None
    return build_rows(entries, site_id=_site_id)


@app.post("/audit-entry-update")
def update_entry(req: AuditEntryUpdateRequest):
    payload = req.dict(exclude_unset=True)
    values = payload.pop("values", {}) or {}
    payload.update(values)

    entry_id = payload.pop("id", None) or payload.pop("audit_id", None)
    batch_id = payload.pop("audit_batch_id", None)
    serial = payload.pop("serial_number", None)

    data = clean_entry_payload(payload)

    boolean_fields = [
        "device_present",
        "camera_sync_off",
        "onedrive_sync_on",
        "password_issue",
        "weak_pin_detected",
        "security_breach",
        "ignore_flag",
    ]
    for field in boolean_fields:
        if field in data:
            data[field] = parse_bool(data[field], default=False if field != "device_present" else True)

    if entry_id:
        result = (
            supabase.table("audit_entries")
            .update(data)
            .eq("audit_id", entry_id)
            .execute()
        )
    elif batch_id and serial:
        result = (
            supabase.table("audit_entries")
            .update(data)
            .eq("audit_batch_id", batch_id)
            .eq("serial_number", serial)
            .execute()
        )
    else:
        raise HTTPException(status_code=400, detail="audit_id or audit_batch_id + serial_number required")

    if not result.data:
        raise HTTPException(status_code=404, detail="Update failed")

    updated_entry = result.data[0]
    invalidate_stored_report(updated_entry.get("audit_batch_id") or batch_id)

    return {"success": True, "entry": updated_entry}


@app.get("/audit-report/{audit_batch_id}")
def get_audit_report(audit_batch_id: str):
    return build_report(audit_batch_id)


@app.post("/audit-report/store/{audit_batch_id}")
def store_audit_report(audit_batch_id: str, force: bool = Query(default=True)):
    path, filename, generated = generate_and_store_report(audit_batch_id, force=force)
    return {
        "success": True,
        "audit_batch_id": audit_batch_id,
        "filename": filename,
        "report_url": path,
        "generated": generated,
        "report_generated": True,
    }


@app.get("/audit-report/export/{audit_batch_id}")
def export_audit_report(audit_batch_id: str):
    path, filename, _ = generate_and_store_report(audit_batch_id, force=True)

    response = FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response




@app.get("/audit-report/executive-pdf/{audit_batch_id}")
def export_executive_pdf(audit_batch_id: str):
    report = build_report(audit_batch_id)
    summary = report.get("summary", {})
    score = calculate_score(summary)
    risk = calculate_risk(summary)

    output_dir = os.getenv(
        "AUDIT_REPORT_STORAGE_DIR",
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_reports"),
    )

    path, filename = generate_executive_pdf(report, score, risk, output_dir)

    response = FileResponse(
        path,
        media_type="application/pdf",
        filename=filename,
    )
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response

@app.get("/director-dashboard")
def director_dashboard():
    # ── 1. All audits — one query ────────────────────────────────────────
    all_audits = (
        supabase.table("device_audits")
        .select("*")
        .order("audit_date", desc=True)
        .execute().data or []
    )
    audits_30 = [
        a for a in all_audits
        if audit_age_days(a.get("audit_date")) is not None
        and audit_age_days(a.get("audit_date")) <= 30
    ]
    batch_ids = [a["audit_batch_id"] for a in audits_30 if a.get("audit_batch_id")]
    batch_map = {a["audit_batch_id"]: a for a in audits_30 if a.get("audit_batch_id")}

    # ── 2. All audit entries — one bulk IN query ─────────────────────────
    if batch_ids:
        all_entries = (
            supabase.table("audit_entries")
            .select("*")
            .in_("audit_batch_id", batch_ids)
            .execute().data or []
        )
    else:
        all_entries = []
    entry_map: dict = {}
    for e in all_entries:
        bid = e.get("audit_batch_id")
        if bid:
            entry_map.setdefault(bid, []).append(e)

    # ── 3. All site configs — one bulk IN query ──────────────────────────
    site_ids = list({a["site_id"] for a in audits_30 if a.get("site_id")})
    if site_ids:
        site_cfgs_raw = (
            supabase.table("sites")
            .select("site_id,photo_retention_days,photo_count_threshold")
            .in_("site_id", site_ids)
            .execute().data or []
        )
    else:
        site_cfgs_raw = []
    cfg_map = {s["site_id"]: s for s in site_cfgs_raw if s.get("site_id")}

    # ── 4. Build summaries in memory — zero additional DB calls ──────────
    def _thresholds(site_id):
        cfg = cfg_map.get(site_id) or {}
        return (
            int(cfg.get("photo_retention_days") or 31),
            int(cfg.get("photo_count_threshold") or 1000),
        )

    def _summarise(batch_id, site_id):
        entries = entry_map.get(batch_id, [])
        ret_days, photo_thr = _thresholds(site_id)
        active = [e for e in entries if not is_true(e.get("ignore_flag"))]
        n_crit = n_warn = n_ok = n_miss = n_photo = n_sec = 0
        for e in active:
            status, _ = classify(e, retention_days=ret_days, photo_threshold=photo_thr)
            if   status == "critical": n_crit += 1
            elif status == "warning":  n_warn += 1
            else:                      n_ok   += 1
            if e.get("device_present") is False:
                n_miss += 1
            age_p = days_old(e.get("photos_date"))
            if is_photo_audit_device(e) and age_p is not None and age_p > ret_days:
                n_photo += 1
            if (e.get("password_issue") is True or
                    e.get("weak_pin_detected") is True or
                    e.get("security_breach") is True):
                n_sec += 1
        return {
            "total_devices":           len(active),
            "ignored":                 len(entries) - len(active),
            "critical":                n_crit,
            "warnings":                n_warn,
            "ok":                      n_ok,
            "missing":                 n_miss,
            "photo_retention_breaches": n_photo,
            "photo_count_warnings":    0,
            "security_issues":         n_sec,
        }

    decorated_reports = []
    for batch_id in batch_ids:
        batch = batch_map.get(batch_id)
        if not batch:
            continue
        site_id = batch.get("site_id")
        summary = _summarise(batch_id, site_id)
        risk    = calculate_risk(summary)
        score   = calculate_score(summary)
        age     = audit_age_days(batch.get("audit_date"))
        decorated_reports.append({
            "site_id":          site_id,
            "site_name":        batch.get("site_name"),
            "audit_batch_id":   batch_id,
            "latest_batch_id":  batch_id,
            "audit_date":       batch.get("audit_date"),
            "audit_age_days":   age,
            "stale_audit":      False,
            "risk_label":       risk,
            "overall_status":   "Compliant" if risk == "On Track" else "Not Compliant",
            "compliance_score": score,
            "report_generated": is_true(batch.get("report_generated")) or bool(batch.get("report_url")),
            "report_url":       batch.get("report_url"),
            "summary":          summary,
        })

    decorated_reports.sort(
        key=lambda r: str(r.get("audit_date") or ""),
        reverse=True,
    )

    latest_by_site: dict = {}
    for report in decorated_reports:
        key = site_key(report.get("site_name")) or site_key(report.get("site_id")) or "unknown"
        if key not in latest_by_site:
            latest_by_site[key] = report
    site_rows = list(latest_by_site.values())

    totals = {
        "sites":                          len(site_rows),
        "audits":                         len(decorated_reports),
        "stored_reports":                 0,
        "total_devices":                  0,
        "critical":                       0,
        "warnings":                       0,
        "ok":                             0,
        "missing":                        0,
        "photo_retention_breaches":       0,
        "security_issues":                0,
        "average_compliance_score":       0,
        "high_risk_sites":                0,
        "needs_attention_sites":          0,
        "on_track_sites":                 0,
        "sites_with_issues_last_30_days": 0,
        "audits_with_issues_last_30_days": 0,
    }
    scores = []
    sites_with_issues: set = set()
    for report in decorated_reports:
        s   = report.get("summary") or {}
        key = site_key(report.get("site_name")) or site_key(report.get("site_id")) or "unknown"
        totals["stored_reports"]             += 1 if report.get("report_generated") else 0
        totals["total_devices"]              += int(s.get("total_devices") or 0)
        totals["critical"]                   += int(s.get("critical")      or 0)
        totals["warnings"]                   += int(s.get("warnings")      or 0)
        totals["ok"]                         += int(s.get("ok")            or 0)
        totals["missing"]                    += int(s.get("missing")       or 0)
        totals["photo_retention_breaches"]   += int(s.get("photo_retention_breaches") or 0)
        totals["security_issues"]            += int(s.get("security_issues") or 0)
        scores.append(int(report.get("compliance_score") or 0))
        if report.get("risk_label") != "On Track":
            totals["audits_with_issues_last_30_days"] += 1
            sites_with_issues.add(key)

    for site in site_rows:
        if   site.get("risk_label") == "High Risk":       totals["high_risk_sites"]      += 1
        elif site.get("risk_label") == "Needs Attention": totals["needs_attention_sites"] += 1
        elif site.get("risk_label") == "On Track":        totals["on_track_sites"]        += 1

    totals["sites_with_issues_last_30_days"] = len(sites_with_issues)
    if scores:
        totals["average_compliance_score"] = round(sum(scores) / len(scores), 1)

    risk_rank = {"High Risk": 0, "Needs Attention": 1, "On Track": 2}
    site_rows.sort(key=lambda r: (
        risk_rank.get(r.get("risk_label"), 9),
        int(r.get("compliance_score") or 0),
        str(r.get("site_name") or ""),
    ))

    return {
        "generated_at":  datetime.now().isoformat(),
        "totals":        totals,
        "sites":         site_rows,
        "recent_audits": decorated_reports,
    }
@app.get("/audit-cleanup/list")
def list_audits_for_cleanup(
    site_id: Optional[str] = Query(default=None),
    site_name: Optional[str] = Query(default=None),
):
    audits = supabase.table("device_audits").select("*").execute().data or []

    resolved_site_name = None

    if site_id:
        try:
            site = resolve_site(site_id=site_id)
            resolved_site_name = site.get("site_name")
        except Exception:
            resolved_site_name = None

    if site_id or site_name:
        wanted_site_id = str(site_id or "").strip()
        wanted_site_name = str(site_name or resolved_site_name or "").strip().lower()

        filtered = []
        for audit in audits:
            audit_site_id = str(audit.get("site_id") or "").strip()
            audit_site_name = str(audit.get("site_name") or "").strip().lower()

            if wanted_site_id and audit_site_id == wanted_site_id:
                filtered.append(audit)
                continue

            if wanted_site_name and audit_site_name == wanted_site_name:
                filtered.append(audit)
                continue

        audits = filtered

    def audit_sort_key(a):
        return str(a.get("audit_date") or a.get("created_at") or "")

    audits = sorted(audits, key=audit_sort_key, reverse=True)

    results = []
    for audit in audits:
        batch_id = audit.get("audit_batch_id")
        entry_count = 0
        if batch_id:
            try:
                entry_rows = (
                    supabase.table("audit_entries")
                    .select("audit_batch_id")
                    .eq("audit_batch_id", batch_id)
                    .execute()
                    .data
                    or []
                )
                entry_count = len(entry_rows)
            except Exception:
                entry_count = 0

        row = dict(audit)
        row["entry_count"] = entry_count
        row["has_report_file"] = bool(audit.get("report_url") and os.path.exists(str(audit.get("report_url"))))
        results.append(row)

    return results


@app.delete("/audit-cleanup/delete/{audit_batch_id}")
def delete_audit_batch(audit_batch_id: str):
    batch = get_audit_batch(audit_batch_id)

    report_url = batch.get("report_url")

    entry_delete = (
        supabase.table("audit_entries")
        .delete()
        .eq("audit_batch_id", audit_batch_id)
        .execute()
    )

    audit_delete = (
        supabase.table("device_audits")
        .delete()
        .eq("audit_batch_id", audit_batch_id)
        .execute()
    )

    file_deleted = False
    if report_url and os.path.exists(str(report_url)):
        try:
            os.remove(str(report_url))
            file_deleted = True
        except Exception:
            file_deleted = False

    return {
        "success": True,
        "audit_batch_id": audit_batch_id,
        "deleted_entries": len(entry_delete.data or []),
        "deleted_audits": len(audit_delete.data or []),
        "file_deleted": file_deleted,
    }


# ── Site Management Endpoints ─────────────────────────────────────────

SITE_FIELDS = {
    "site_name", "company_name", "location", "address",
    "contact_name", "contact_email", "contact_phone",
    "notes", "active",
    "sharepoint_tenant_id", "sharepoint_client_id", "sharepoint_client_secret",
    "sharepoint_site_url", "sharepoint_drive_id", "sharepoint_folder",
    "sharepoint_enabled",
    # Compliance policy thresholds
    "photo_retention_days", "photo_count_threshold", "audit_frequency_days",
    # Policy checklist booleans
    "consent_forms_current", "photo_policy_documented", "staff_training_current",
    "encryption_enforced", "access_controls_documented", "data_retention_policy",
    "incident_response_plan", "privacy_impact_assessment", "third_party_data_agreement",
}


def clean_site_payload(payload: dict) -> dict:
    return {k: v for k, v in payload.items() if k in SITE_FIELDS}


@app.get("/sites/{site_id}")
def get_site_detail(site_id: str):
    """Full site detail including SharePoint config."""
    r = supabase.table("sites").select("*").eq("site_id", site_id).limit(1).execute()
    if not r.data:
        raise HTTPException(status_code=404, detail="Site not found")
    return r.data[0]


@app.post("/sites")
def create_site(payload: Dict[str, Any]):
    """Create a new site."""
    site_name = payload.get("site_name", "").strip()
    if not site_name:
        raise HTTPException(status_code=400, detail="site_name is required")

    existing = (
        supabase.table("sites")
        .select("site_id, site_name")
        .eq("site_name", site_name)
        .limit(1)
        .execute()
    )
    if existing.data:
        raise HTTPException(
            status_code=409,
            detail=f"Site '{site_name}' already exists. Select it from the dropdown instead.",
        )

    data = clean_site_payload(payload)
    data["site_name"] = site_name
    data.setdefault("active", True)
    data.setdefault("sharepoint_enabled", False)

    result = supabase.table("sites").insert(data).execute()
    if not result.data:
        raise HTTPException(status_code=500, detail="Failed to create site")

    return result.data[0]


@app.put("/sites/{site_id}")
def update_site(site_id: str, payload: Dict[str, Any]):
    """Update site fields (contact info, SharePoint config, etc.)."""
    data = clean_site_payload(payload)

    if not data:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    result = (
        supabase.table("sites")
        .update(data)
        .eq("site_id", site_id)
        .execute()
    )

    if not result.data:
        raise HTTPException(status_code=404, detail="Site not found or update failed")

    return {"success": True, "site": result.data[0]}


@app.post("/sites/{site_id}/test-sharepoint")
def test_site_sharepoint(site_id: str):
    """Test SharePoint connection using stored credentials for a site."""
    if not SHAREPOINT_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="SharePoint integration not available — msal package not installed",
        )

    r = supabase.table("sites").select("*").eq("site_id", site_id).limit(1).execute()
    if not r.data:
        raise HTTPException(status_code=404, detail="Site not found")

    site = r.data[0]
    return sp_test_connection(site)


# ── SharePoint Upload Endpoint ────────────────────────────────────────

@app.post("/audit-report/upload-sharepoint/{audit_batch_id}")
def upload_report_sharepoint(audit_batch_id: str):
    """Upload the PDF report for a batch to the site's configured SharePoint."""
    if not SHAREPOINT_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="SharePoint integration not available — msal package not installed",
        )

    batch = get_audit_batch(audit_batch_id)
    site_id = batch.get("site_id")
    if not site_id:
        raise HTTPException(status_code=400, detail="Audit batch has no site_id")

    site_r = supabase.table("sites").select("*").eq("site_id", site_id).limit(1).execute()
    if not site_r.data:
        raise HTTPException(status_code=404, detail="Site not found")

    site = site_r.data[0]

    if not site.get("sharepoint_enabled"):
        raise HTTPException(status_code=400, detail="SharePoint upload not enabled for this site")

    # Generate the PDF first
    report = build_report(audit_batch_id)
    summary = report.get("summary", {})
    score = calculate_score(summary)
    risk = calculate_risk(summary)

    output_dir = os.getenv(
        "AUDIT_REPORT_STORAGE_DIR",
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_reports"),
    )

    pdf_path, pdf_filename = generate_executive_pdf(report, score, risk, output_dir)

    # Upload to SharePoint
    try:
        result = upload_report_to_sharepoint(site, pdf_path, pdf_filename)
    except Exception as e:
        logger.error(f"SharePoint upload failed for batch {audit_batch_id}: {e}")
        raise HTTPException(status_code=502, detail=f"SharePoint upload failed: {e}")

    # Store the SharePoint URL in the batch metadata
    try:
        supabase.table("device_audits").update({
            "report_generated": True,
            "report_url": pdf_path,
        }).eq("audit_batch_id", audit_batch_id).execute()
    except Exception:
        pass  # Non-critical — report was already uploaded

    return {
        "success": True,
        "audit_batch_id": audit_batch_id,
        "sharepoint_url": result.get("sharepoint_url"),
        "filename": result.get("filename"),
        "size": result.get("size"),
    }


# ── Remediation Tracking Endpoints ────────────────────────────────────

REMEDIATION_FIELDS = {
    "site_id", "audit_batch_id", "source", "finding_type", "device_serial",
    "description", "assigned_to", "priority", "status", "due_date",
    "resolved_date", "resolution_notes",
}


def clean_remediation_payload(payload: dict) -> dict:
    return {k: v for k, v in payload.items() if k in REMEDIATION_FIELDS}


@app.get("/remediation-actions")
def list_remediation_actions(
    site_id: Optional[str] = Query(default=None),
    status: Optional[str] = Query(default=None),
    audit_batch_id: Optional[str] = Query(default=None),
    priority: Optional[str] = Query(default=None),
):
    q = supabase.table("remediation_actions").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    if status:
        q = q.eq("status", status)
    if audit_batch_id:
        q = q.eq("audit_batch_id", audit_batch_id)
    if priority:
        q = q.eq("priority", priority)
    try:
        q = q.order("created_at", desc=True)
    except TypeError:
        q = q.order("created_at")
    return q.execute().data or []


@app.post("/remediation-actions")
def create_remediation_action(payload: Dict[str, Any]):
    data = clean_remediation_payload(payload)
    if not data.get("description"):
        raise HTTPException(status_code=400, detail="description is required")
    if not data.get("site_id"):
        raise HTTPException(status_code=400, detail="site_id is required")
    data.setdefault("source", "manual")
    data.setdefault("priority", "medium")
    data.setdefault("status", "open")
    result = supabase.table("remediation_actions").insert(data).execute()
    if not result.data:
        raise HTTPException(status_code=500, detail="Failed to create action")
    return result.data[0]


@app.put("/remediation-actions/{action_id}")
def update_remediation_action(action_id: str, payload: Dict[str, Any]):
    data = clean_remediation_payload(payload)
    if not data:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    # Auto-set resolved_date when status changes to resolved
    if data.get("status") in ("resolved", "accepted_risk") and not data.get("resolved_date"):
        data["resolved_date"] = date.today().isoformat()

    # Clear resolved_date if re-opened
    if data.get("status") in ("open", "in_progress"):
        data["resolved_date"] = None

    data["updated_at"] = datetime.now().isoformat()

    result = (
        supabase.table("remediation_actions")
        .update(data)
        .eq("action_id", action_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Action not found")
    return {"success": True, "action": result.data[0]}


@app.post("/remediation-actions/generate/{audit_batch_id}")
def generate_remediation_actions(audit_batch_id: str):
    """Auto-generate remediation actions from audit findings."""
    batch = get_audit_batch(audit_batch_id)
    site_id = batch.get("site_id")
    rows = get_entries_for_batch(audit_batch_id)
    active_rows = [r for r in rows if not is_true(r.get("ignore_flag"))]

    # Check for existing actions for this batch to avoid duplicates
    existing = (
        supabase.table("remediation_actions")
        .select("device_serial, finding_type")
        .eq("audit_batch_id", audit_batch_id)
        .execute()
        .data or []
    )
    existing_keys = set()
    for e in existing:
        existing_keys.add((e.get("device_serial"), e.get("finding_type")))

    actions_to_create = []
    default_due = (date.today() + __import__("datetime").timedelta(days=14)).isoformat()

    for row in active_rows:
        serial = row.get("serial_number")
        rm = room_name(row)
        model = device_label(row)
        dtype = row.get("device_type") or "Device"

        # Missing device
        if is_missing(row) and (serial, "missing_device") not in existing_keys:
            actions_to_create.append({
                "site_id": site_id,
                "audit_batch_id": audit_batch_id,
                "source": "audit",
                "finding_type": "missing_device",
                "device_serial": serial,
                "description": f"{dtype} {model} ({rm}) - missing from audit. Locate, verify and disable access if unrecovered.",
                "priority": "critical",
                "status": "open",
                "due_date": default_due,
            })

        # Photo retention
        if is_photo_action(row) and (serial, "photo_retention") not in existing_keys:
            pc = photo_count(row)
            pd = row.get("photos_date")
            actions_to_create.append({
                "site_id": site_id,
                "audit_batch_id": audit_batch_id,
                "source": "audit",
                "finding_type": "photo_retention",
                "device_serial": serial,
                "description": f"{dtype} {model} ({rm}) - {pc} photos, date {format_report_value(pd)}. Reduce to under 30 days.",
                "priority": "high",
                "status": "open",
                "due_date": default_due,
            })

        # Credential issues
        if has_credential_issue(row) and (serial, "credential") not in existing_keys:
            notes = normalize_text(row.get("breach_notes") or row.get("notes"))
            actions_to_create.append({
                "site_id": site_id,
                "audit_batch_id": audit_batch_id,
                "source": "audit",
                "finding_type": "credential",
                "device_serial": serial,
                "description": f"{dtype} {model} ({rm}) - credential/PIN issue. {notes}. Apply unique device credentials.",
                "priority": "critical",
                "status": "open",
                "due_date": default_due,
            })

        # Hardware issues
        if has_hardware_issue(row) and (serial, "hardware") not in existing_keys:
            notes = normalize_text(row.get("breach_notes") or row.get("notes"))
            actions_to_create.append({
                "site_id": site_id,
                "audit_batch_id": audit_batch_id,
                "source": "audit",
                "finding_type": "hardware",
                "device_serial": serial,
                "description": f"{dtype} {model} ({rm}) - hardware issue. {notes}. Repair, replace or risk-accept.",
                "priority": "high",
                "status": "open",
                "due_date": default_due,
            })

    created_count = 0
    if actions_to_create:
        result = supabase.table("remediation_actions").insert(actions_to_create).execute()
        created_count = len(result.data or [])

    return {
        "success": True,
        "audit_batch_id": audit_batch_id,
        "site_id": site_id,
        "actions_created": created_count,
        "actions_skipped_duplicate": len(active_rows) - created_count,
    }

@app.get("/governance-summary")
def governance_summary():
    """Aggregate remediation, incident and policy data with NQS quality area mapping."""
    from datetime import date as _date

    today = str(_date.today())

    # ── Remediation ──────────────────────────────────────────────────────────
    try:
        rem = supabase.table("remediation_actions").select(
            "status,priority,due_date,finding_type,site_name"
        ).execute().data or []
    except Exception:
        rem = []

    rem_open = [r for r in rem if r.get("status") in ("open", "in_progress")]
    rem_overdue = [
        r for r in rem_open
        if r.get("due_date") and str(r.get("due_date")) < today
    ]
    rem_resolved = [r for r in rem if r.get("status") in ("resolved", "closed")]

    # ── Incidents ─────────────────────────────────────────────────────────────
    try:
        inc = supabase.table("incidents").select(
            "status,severity,mandatory_reportable,incident_type"
        ).execute().data or []
    except Exception:
        inc = []

    inc_open = [i for i in inc if i.get("status") not in ("closed", "resolved")]
    inc_mandatory = [i for i in inc if i.get("mandatory_reportable") is True]

    # ── Policy ────────────────────────────────────────────────────────────────
    policy_keys = [
        "consent_forms_current", "photo_policy_documented", "staff_training_current",
        "encryption_enforced", "access_controls_documented", "data_retention_policy",
        "incident_response_plan", "privacy_impact_assessment", "third_party_data_agreement",
    ]
    try:
        sites = supabase.table("sites").select(
            "site_id,site_name," + ",".join(policy_keys)
        ).execute().data or []
    except Exception:
        sites = []

    policy_labels = {
        "consent_forms_current":     "Consent forms current",
        "photo_policy_documented":   "Photo policy documented",
        "staff_training_current":    "Staff privacy training current",
        "encryption_enforced":       "Device encryption enforced",
        "access_controls_documented":"Access controls documented",
        "data_retention_policy":     "Data retention policy in place",
        "incident_response_plan":    "Incident response plan exists",
        "privacy_impact_assessment": "Privacy impact assessment completed",
        "third_party_data_agreement":"Third-party data agreements in place",
    }
    policy_gaps = 0
    policy_complete = 0
    policy_by_site = []
    for site in sites:
        site_bools = [site.get(k) for k in policy_keys]
        missing_labels = [policy_labels[k] for k in policy_keys if site.get(k) is False]
        not_set_labels = [policy_labels[k] for k in policy_keys if site.get(k) is None]
        confirmed = sum(1 for v in site_bools if v is True)
        gaps = sum(1 for v in site_bools if v is False)
        if all(v is True for v in site_bools):
            policy_complete += 1
        else:
            policy_gaps += gaps
        policy_by_site.append({
            "site_id":    site.get("site_id"),
            "site_name":  site.get("site_name") or site.get("site_id"),
            "confirmed":  confirmed,
            "gaps":       gaps,
            "not_set":    len(not_set_labels),
            "missing":    missing_labels,
            "incomplete": not_set_labels,
        })
    policy_by_site.sort(key=lambda s: s["gaps"], reverse=True)

    # ── NQS Quality Area Mapping ──────────────────────────────────────────────
    def count_rem(keywords):
        return sum(
            1 for r in rem_open
            if any(kw in (r.get("finding_type") or "").lower() for kw in keywords)
        )

    nqs_mapping = [
        {
            "area_code": "NQS 2.2",
            "area_name": "Children's Health & Safety",
            "description": "Photo retention compliance, incident management and mandatory notification obligations.",
            "open_remediations": count_rem(["photo", "photo_retention"]),
            "open_incidents": len(inc_open),
            "mandatory_incidents": len(inc_mandatory),
            "policy_gaps": 0,
        },
        {
            "area_code": "NQS 3.2",
            "area_name": "Use of Space and Equipment",
            "description": "Missing or unaccounted devices, hardware condition and asset register integrity.",
            "open_remediations": count_rem(["hardware", "missing", "device"]),
            "open_incidents": 0,
            "mandatory_incidents": 0,
            "policy_gaps": 0,
        },
        {
            "area_code": "NQS 4.2",
            "area_name": "Professionalism",
            "description": "Staff data handling, privacy training currency and credential management.",
            "open_remediations": count_rem(["credential", "security", "password", "pin"]),
            "open_incidents": 0,
            "mandatory_incidents": 0,
            "policy_gaps": sum(
                1 for s in sites
                if s.get("staff_training_current") is False
            ),
        },
        {
            "area_code": "NQS 7.3",
            "area_name": "Governance — Roles & Responsibilities",
            "description": "Data retention policies, access controls, privacy documentation and third-party agreements.",
            "open_remediations": len(rem_open),
            "open_incidents": 0,
            "mandatory_incidents": 0,
            "policy_gaps": policy_gaps,
        },
    ]

    # ── Audit recency ─────────────────────────────────────────────────────────
    try:
        recent_audits = (
            supabase.table("device_audits")
            .select("site_id,site_name,audit_date")
            .order("audit_date", desc=True)
            .execute().data or []
        )
        # Latest audit per site
        seen = {}
        for a in recent_audits:
            sid = a.get("site_id")
            if sid and sid not in seen:
                seen[sid] = a
        from datetime import date as _d2
        today_dt = _d2.today()
        audit_recency = []
        for sid, a in seen.items():
            try:
                adate = _d2.fromisoformat(str(a.get("audit_date",""))[:10])
                age = (today_dt - adate).days
            except Exception:
                age = None
            audit_recency.append({
                "site_id":    sid,
                "site_name":  a.get("site_name") or sid,
                "audit_date": a.get("audit_date"),
                "age_days":   age,
                "overdue":    age is not None and age > 90,
            })
        audit_recency.sort(key=lambda x: (x.get("age_days") or 9999), reverse=True)
    except Exception:
        audit_recency = []

    return {
        "generated_at": today,
        "remediation": {
            "total": len(rem),
            "open": len(rem_open),
            "overdue": len(rem_overdue),
            "resolved": len(rem_resolved),
            "by_priority": {
                "critical": sum(1 for r in rem if r.get("priority") == "critical"),
                "high": sum(1 for r in rem if r.get("priority") == "high"),
                "medium": sum(1 for r in rem if r.get("priority") == "medium"),
                "low": sum(1 for r in rem if r.get("priority") == "low"),
            },
        },
        "incidents": {
            "total": len(inc),
            "open": len(inc_open),
            "mandatory_reportable": len(inc_mandatory),
            "by_severity": {
                "critical": sum(1 for i in inc if i.get("severity") == "critical"),
                "high": sum(1 for i in inc if i.get("severity") == "high"),
                "medium": sum(1 for i in inc if i.get("severity") == "medium"),
                "low": sum(1 for i in inc if i.get("severity") == "low"),
            },
        },
        "policy": {
            "sites_total": len(sites),
            "policy_complete_sites": policy_complete,
            "policy_gaps": policy_gaps,
            "by_site": policy_by_site,
        },
        "audit_recency": audit_recency,
        "nqs_mapping": nqs_mapping,
    }




@app.get("/remediation-summary")
def remediation_summary(site_id: Optional[str] = Query(default=None)):
    q = supabase.table("remediation_actions").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    actions = q.execute().data or []

    today = date.today().isoformat()
    summary = {"open": 0, "in_progress": 0, "overdue": 0, "resolved": 0, "accepted_risk": 0, "total": len(actions)}

    for a in actions:
        s = a.get("status", "open")
        if s in summary:
            summary[s] += 1
        if s in ("open", "in_progress") and a.get("due_date") and a["due_date"] < today:
            summary["overdue"] += 1

    return summary


# ── Incident Register Endpoints ───────────────────────────────────────

INCIDENT_FIELDS = {
    "site_id", "incident_date", "reported_by", "incident_type", "severity",
    "description", "children_affected", "affected_count", "mandatory_reportable",
    "reported_to", "reported_date", "response_actions", "status",
    "resolved_date", "resolution_notes", "linked_audit_batch_id",
    "linked_device_serial",
}


def clean_incident_payload(payload: dict) -> dict:
    return {k: v for k, v in payload.items() if k in INCIDENT_FIELDS}


@app.get("/incidents")
def list_incidents(
    site_id: Optional[str] = Query(default=None),
    status: Optional[str] = Query(default=None),
    severity: Optional[str] = Query(default=None),
):
    q = supabase.table("incidents").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    if status:
        q = q.eq("status", status)
    if severity:
        q = q.eq("severity", severity)
    try:
        q = q.order("created_at", desc=True)
    except TypeError:
        q = q.order("created_at")
    return q.execute().data or []


@app.post("/incidents")
def create_incident(payload: Dict[str, Any]):
    data = clean_incident_payload(payload)
    if not data.get("description"):
        raise HTTPException(status_code=400, detail="description is required")
    if not data.get("site_id"):
        raise HTTPException(status_code=400, detail="site_id is required")
    data.setdefault("incident_type", "other")
    data.setdefault("severity", "medium")
    data.setdefault("status", "open")
    data.setdefault("incident_date", date.today().isoformat())
    data.setdefault("children_affected", False)
    data.setdefault("mandatory_reportable", False)

    # Convert boolean strings
    for field in ("children_affected", "mandatory_reportable"):
        if field in data:
            data[field] = parse_bool(data[field], default=False)

    result = supabase.table("incidents").insert(data).execute()
    if not result.data:
        raise HTTPException(status_code=500, detail="Failed to create incident")
    return result.data[0]


@app.put("/incidents/{incident_id}")
def update_incident(incident_id: str, payload: Dict[str, Any]):
    data = clean_incident_payload(payload)
    if not data:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    # Auto-set resolved_date
    if data.get("status") in ("resolved", "closed") and not data.get("resolved_date"):
        data["resolved_date"] = date.today().isoformat()
    if data.get("status") in ("open", "investigating"):
        data["resolved_date"] = None

    # Convert boolean strings
    for field in ("children_affected", "mandatory_reportable"):
        if field in data:
            data[field] = parse_bool(data[field], default=False)

    data["updated_at"] = datetime.now().isoformat()

    result = (
        supabase.table("incidents")
        .update(data)
        .eq("incident_id", incident_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Incident not found")
    return {"success": True, "incident": result.data[0]}


@app.delete("/incidents/{incident_id}")
def delete_incident(incident_id: str):
    result = (
        supabase.table("incidents")
        .delete()
        .eq("incident_id", incident_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Incident not found")
    return {"success": True}


@app.get("/incident-summary")
def incident_summary(site_id: Optional[str] = Query(default=None)):
    q = supabase.table("incidents").select("*")
    if site_id:
        q = q.eq("site_id", site_id)
    incidents = q.execute().data or []

    summary = {
        "open": 0, "investigating": 0, "reported": 0,
        "resolved": 0, "closed": 0, "total": len(incidents),
        "mandatory_reportable": 0, "children_affected": 0,
    }

    for i in incidents:
        s = i.get("status", "open")
        if s in summary:
            summary[s] += 1
        if i.get("mandatory_reportable"):
            summary["mandatory_reportable"] += 1
        if i.get("children_affected"):
            summary["children_affected"] += 1

    return summary
